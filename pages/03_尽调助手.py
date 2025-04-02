import streamlit as st
import pandas as pd
import datetime
import time
from docx import Document
import tempfile

# ========== 页面基础设置 ==========
st.set_page_config(page_title="尽调助手（整合版）", layout="wide")
st.markdown("📍 当前流程节点：尽调助手（第 3 步 / 共 4 步）")

# ======= 顶部：上下文客户显示 =======
client = st.session_state.get("active_client")
st.markdown("## 🧭 整合式尽调助手")
if client:
    st.markdown(f"**当前尽调客户：** `{client}`")
else:
    st.info("当前暂无已选客户，可从客户画像模块点击 '进入尽调' 传递上下文。")

st.markdown("---")

# ========== 左侧导航栏 ==============
menu = st.sidebar.radio(
    "选择尽调功能",
    ["任务管理", "问题生成", "建议分析", "模拟对话", "报告预览与导出"]
)

# ========== 数据初始化（session_state） ==========

# 1) 任务管理数据
if "task_data" not in st.session_state:
    st.session_state.task_data = [
        {"企业": "宇树科技", "负责人": "张三", "状态": "进行中", 
         "启动日期": datetime.date(2025, 3, 1), 
         "预计完成": datetime.date(2025, 3, 10)},
        {"企业": "优必选", "负责人": "李四", "状态": "已完成", 
         "启动日期": datetime.date(2025, 2, 15), 
         "预计完成": datetime.date(2025, 2, 25)},
        {"企业": "智元机器人", "负责人": "王五", "状态": "待启动", 
         "启动日期": datetime.date(2025, 3, 20), 
         "预计完成": datetime.date(2025, 3, 30)},
    ]

# 2) 问题生成数据
if "question_data" not in st.session_state:
    st.session_state.question_data = {
        "宏观": [], "中观": [], "微观": []
    }

# 默认模板（同原02_尽调问题生成）
default_questions = {
    "宏观": [
        "公司如何解读国家最新机器人产业支持政策？",
        "当前国内经济环境对贵公司主要业务有哪些影响？",
        "国际贸易形势变化是否影响原材料成本？",
        "公司对近期高层讲话中关于智能制造的表述有何回应？",
        "贵公司是否参与地方产业规划或试点计划？",
    ],
    "中观": [
        "公司与哪些核心供应商/客户保持合作？",
        "当前行业竞争格局有哪些新变化？",
        "是否存在上下游集中度过高的风险？",
        "公司在行业中的市场份额及定位？",
        "未来6-12个月产品或市场有何调整规划？",
    ],
    "微观": [
        "企业是否存在短期到期债务？资金如何安排？",
        "公司过去一年净利润与毛利率趋势如何？",
        "研发投入占比及未来是否继续增加？",
        "团队人员结构是否有明显调整计划？",
        "主要资金需求用途是什么？对应回报预期如何？",
    ]
}
for lvl in ["宏观","中观","微观"]:
    if not st.session_state.question_data[lvl]:
        st.session_state.question_data[lvl] = [
            {"question": q, "answer": "", "include": True} 
            for q in default_questions[lvl]
        ]

# 3) 报告问题同步
if "report_questions" not in st.session_state:
    st.session_state.report_questions = {}

# ========== 定义一些通用函数 ============

def _refresh_page():
    # 简易刷新机制
    st.experimental_rerun()

def generate_word_report(company, manager, date_val, question_data, summary_text):
    doc = Document()
    doc.add_heading(f'客户尽调报告 - {company}', 0)
    doc.add_paragraph(f"客户经理：{manager}")
    doc.add_paragraph(f"报告日期：{date_val.strftime('%Y-%m-%d')}")
    doc.add_paragraph("")
    for section, title in zip(["宏观", "中观", "微观"], ["一、宏观层面", "二、中观层面", "三、微观层面"]):
        doc.add_heading(title, level=1)
        for idx, item in enumerate(question_data.get(section, []), 1):
            doc.add_paragraph(f"{idx}. {item['question']}", style='List Number')
            doc.add_paragraph(f"答：{item['answer']}", style='Normal')
    doc.add_heading("四、尽调结论与建议", level=1)
    doc.add_paragraph(summary_text or "暂无")
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        doc.save(tmp.name)
        return tmp.name

# ========== 根据选单menu展示对应模块 ==========

if menu == "任务管理":
    st.markdown("### ① 尽调任务管理")

    # 同01_尽调任务管理逻辑
    # ========== 筛选器 =============
    keyword = st.sidebar.text_input("关键字（企业/负责人）")
    all_tasks = st.session_state.task_data
    active_tasks = [t for t in all_tasks if t["状态"] in ["待启动", "进行中"]]
    completed_tasks = [t for t in all_tasks if t["状态"] == "已完成"]
    if keyword:
        active_tasks = [t for t in active_tasks if keyword in t["企业"] or keyword in t["负责人"]]
        completed_tasks = [t for t in completed_tasks if keyword in t["企业"] or keyword in t["负责人"]]

    st.markdown("#### 📋 当前任务列表")
    if active_tasks:
        for i, task in enumerate(active_tasks):
            cols = st.columns([5, 1, 1, 1])
            with cols[0]:
                st.markdown(f"**🏢 {task['企业']}** — 负责人：{task['负责人']} — 状态：{task['状态']}  ")
                st.markdown(f"⏱️ {task['启动日期'].strftime('%Y-%m-%d')} ~ {task['预计完成'].strftime('%Y-%m-%d')}")
            with cols[1]:
                if st.button("✔ 完成", key=f"complete_{i}"):
                    idx = st.session_state.task_data.index(task)
                    st.session_state.task_data[idx]["状态"] = "已完成"
                    st.success("任务已标记为完成")
                    _refresh_page()
            with cols[2]:
                if st.button("✏ 修改", key=f"edit_{i}"):
                    st.session_state.edit_task_index = st.session_state.task_data.index(task)
                    _refresh_page()
            with cols[3]:
                if st.button("🗑 删除", key=f"delete_{i}"):
                    st.session_state.task_data.remove(task)
                    st.success("任务已删除")
                    _refresh_page()
    else:
        st.info("暂无当前任务")

    st.markdown("---")
    st.markdown("#### ✅ 已完成任务列表")
    if completed_tasks:
        for i, task in enumerate(completed_tasks):
            cols = st.columns([5, 1, 1])
            with cols[0]:
                st.markdown(f"**🏢 {task['企业']}** — 负责人：{task['负责人']}  ")
                st.markdown(f"📅 完成时间：{task['预计完成'].strftime('%Y-%m-%d')}")
            with cols[1]:
                if st.button("🔁 恢复", key=f"restore_{i}"):
                    idx = st.session_state.task_data.index(task)
                    st.session_state.task_data[idx]["状态"] = "进行中"
                    st.success("任务已恢复")
                    _refresh_page()
            with cols[2]:
                if st.button("🗑 删除完", key=f"del_done_{i}"):
                    st.session_state.task_data.remove(task)
                    st.success("已完成任务已删除")
                    _refresh_page()
    else:
        st.info("暂无已完成任务")

    # 编辑任务
    if "edit_task_index" in st.session_state:
        idx = st.session_state.edit_task_index
        task = st.session_state.task_data[idx]
        st.markdown("---")
        st.subheader("📝 编辑任务")
        with st.form("edit_task_form"):
            new_name = st.text_input("企业名称", value=task["企业"])
            new_manager = st.text_input("负责人", value=task["负责人"])
            new_status = st.selectbox("状态", ["待启动", "进行中", "已完成"], 
                                      index=["待启动","进行中","已完成"].index(task["状态"]))
            new_start = st.date_input("启动日期", value=task["启动日期"])
            new_end = st.date_input("预计完成日期", value=task["预计完成"])
            save = st.form_submit_button("💾 保存修改")
        if save:
            st.session_state.task_data[idx] = {
                "企业": new_name.strip(),
                "负责人": new_manager.strip(),
                "状态": new_status,
                "启动日期": new_start,
                "预计完成": new_end
            }
            del st.session_state.edit_task_index
            st.success("✅ 修改成功")
            _refresh_page()

    # 添加新任务
    st.markdown("---")
    st.subheader("🆕 创建新尽调任务")
    with st.form("new_task_form"):
        company = st.text_input("企业名称")
        manager = st.text_input("负责人")
        status = st.selectbox("当前状态", ["待启动", "进行中", "已完成"])
        start_date = st.date_input("启动日期", value=datetime.date.today())
        end_date = st.date_input("预计完成日期")
        submitted = st.form_submit_button("📌 创建任务")
    if submitted:
        if company.strip() and manager.strip():
            st.session_state.task_data.append({
                "企业": company.strip(),
                "负责人": manager.strip(),
                "状态": status,
                "启动日期": start_date,
                "预计完成": end_date
            })
            st.success(f"✅ 新任务已创建：{company}（负责人：{manager}）")
            _refresh_page()
        else:
            st.warning("⚠️ 企业名称和负责人不能为空。")

elif menu == "问题生成":
    st.markdown("### ② 尽调问题生成（宏观 / 中观 / 微观）")

    # 引用02_尽调问题生成逻辑
    #--------------------------------------
    if "question_data" not in st.session_state:
        st.session_state.question_data = {"宏观": [], "中观": [], "微观": []}

    def render_question_block(level_key):
        st.markdown(f"#### 🧩 {level_key} 层级问题")

        # 表单新增问题
        with st.form(key=f"form_add_{level_key}"):
            new_q = st.text_input(f"➕ 新增{level_key}层级问题", key=f"input_{level_key}")
            submitted = st.form_submit_button("添加问题")
            if submitted:
                if new_q.strip():
                    st.session_state.question_data[level_key].append({
                        "question": new_q.strip(), 
                        "answer": "", 
                        "include": False
                    })
                    st.success("✅ 已添加新问题")
                    time.sleep(0.5)
                else:
                    st.warning("⚠️ 问题内容不能为空。")

        # 渲染问题列表
        questions = st.session_state.question_data[level_key]
        new_list = []
        for i, q in enumerate(questions):
            cols = st.columns([5, 3, 1, 1])
            with cols[0]:
                q_text = st.text_input(f"{level_key}问题 {i+1}", value=q["question"], key=f"{level_key}_q_{i}")
            with cols[1]:
                a_text = st.text_area(f"回答 {i+1}", value=q["answer"], key=f"{level_key}_a_{i}")
            with cols[2]:
                mark = st.checkbox("纳入报告", value=q.get("include", False), key=f"{level_key}_inc_{i}")
            with cols[3]:
                remove = st.checkbox("删除", key=f"{level_key}_del_{i}")
            if not remove:
                new_list.append({"question": q_text, "answer": a_text, "include": mark})
        st.session_state.question_data[level_key] = new_list

    with st.expander("🌐 宏观层面（政策、行业趋势、经济环境）", expanded=True):
        render_question_block("宏观")
    with st.expander("🏭 中观层面（产业链、市场格局、技术方向）", expanded=True):
        render_question_block("中观")
    with st.expander("🧬 微观层面（财务、团队、产品、融资）", expanded=True):
        render_question_block("微观")

    st.markdown("---")
    if st.button("📤 保存并同步至报告草稿"):
        st.session_state.report_questions = {
            k: [q for q in lst if q["include"]] 
            for k, lst in st.session_state.question_data.items()
        }
        st.success("✅ 问题与回答已保存，可前往报告页查看。")


elif menu == "建议分析":
    st.markdown("### ③ 尽调建议分析")
    st.markdown("该模块提供两部分内容：\n- **外部情报建议**：根据产业趋势、政策舆情\n- **内部补充信息分析**：结合企业回答、财务等")

    # 外部情报
    external_insights = [
        {
            "建议": "行业同类企业 Figure AI 推出新品，贵公司技术路线是否差异化？",
            "来源": "研报《2025人形机器人蓝皮书》3章",
            "操作建议": "追问产品差异与壁垒"
        },
        {
            "建议": "国家对智能制造推出产业引导政策，企业是否已享受？",
            "来源": "工信部政策汇总",
            "操作建议": "确认政策红利"
        }
    ]
    st.subheader("📡 外部风险提醒")
    for item in external_insights:
        st.markdown(f"🔶 **建议：** {item['建议']}")
        st.markdown(f"📎 来源：{item['来源']}")
        st.markdown(f"🧭 操作建议：{item['操作建议']}")
        st.markdown("---")

    # 内部分析
    st.subheader("📊 内部补充信息分析")
    data = st.session_state.get("report_questions", {})
    internal_findings = []
    for sec in ["宏观","中观","微观"]:
        for q in data.get(sec, []):
            if q["question"] and q["answer"]:
                txt = q["question"] + q["answer"]
                if "现金流" in txt or "偿债" in txt:
                    internal_findings.append({
                        "风险点": "企业现金流承压",
                        "来源问题": q["question"]
                    })
                if "债务" in txt and "到期" in txt:
                    internal_findings.append({
                        "风险点": "短期债务到期压力",
                        "来源问题": q["question"]
                    })
    if internal_findings:
        for item in internal_findings:
            st.markdown(f"🔴 **{item['风险点']}**")
            st.markdown(f"📎 来源：{item['来源问题']}")
    else:
        st.info("暂无自动识别到的内部风险项")

    st.markdown("---")
    st.subheader("📥 补充信息 + 模拟建议")
    with st.form("extra_info_form"):
        cash_flow = st.number_input("季度现金流（万）", value=-500)
        top_client_ratio = st.slider("头部客户占比%", 0,100,60)
        ret_days = st.number_input("回款周期（天）", value=85)
        debt_due = st.number_input("近3月到期债务（万）", value=1500)
        sub_btn = st.form_submit_button("分析")
    if sub_btn:
        st.markdown("### 小模型风险建议（模拟）")
        if cash_flow<0:
            st.warning("🔴 现金流净额为负，需关注短期偿债与资金用途")
        if top_client_ratio>50:
            st.warning("🟠 客户集中度高，存在依赖单一大客户风险")
        if ret_days>60:
            st.info("🟡 回款周期偏长或需垫资")
        if debt_due>1000:
            st.warning("🔴 短期债务高额，需要融资安排")
        if all([cash_flow>=0,top_client_ratio<=50,ret_days<=60,debt_due<=1000]):
            st.success("✅ 补充信息无明显风险项")

elif menu == "模拟对话":
    st.markdown("### ④ 尽调模拟对话")

    if "chat_history" not in st.session_state:
        st.session_state.chat_history = []

    user_input = st.text_input("客户经理输入内容：","")
    if user_input:
        st.session_state.chat_history.append(("客户经理",user_input))
        # 模拟回答
        mock_resp = "这是企业的回复示例，强调公司具备灵巧手差异化。"
        st.session_state.chat_history.append(("企业代表", mock_resp))

    for role,msg in st.session_state.chat_history:
        if role=="客户经理":
            st.chat_message("user").markdown(f"**客户经理：** {msg}")
        else:
            st.chat_message("assistant").markdown(f"**企业代表：** {msg}")

    if st.button("清空对话"):
        st.session_state.chat_history=[]

elif menu == "报告预览与导出":
    st.markdown("### ⑤ 报告预览与导出")

    # 类似05_报告预览与导出
    company = st.text_input("企业名称", value=client or "未指定企业")
    manager = st.text_input("客户经理", value="张三")
    report_date = st.date_input("生成日期", value=datetime.date.today())
    data = st.session_state.get("report_questions", {})
    st.markdown("---")
    st.subheader("🧠 总体结论与建议")
    summary = st.text_area("请输入尽调结论与合作建议", height=200)
    st.markdown("---")
    st.subheader("📥 下载Word报告")
    if st.button("生成并下载"):
        word_path = generate_word_report(company, manager, report_date, data, summary)
        with open(word_path, "rb") as f:
            st.download_button("📄 下载报告.docx", f, file_name=f"{company}_尽调报告.docx")

    st.caption("✅ 导出结果可直接在Word中编辑、打印。")

# ============= END =============
st.markdown("---")
st.caption("© 整合版尽调助手 · Streamlit单页模式 · Demo")
